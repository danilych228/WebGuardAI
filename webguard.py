#!/usr/bin/env python3
"""
WebGuardAI - Advanced Code Security Scanner
Сканер безопасности кода на основе AI
"""

import sys
import os
import asyncio
import argparse
import pyfiglet
import requests
import json
import smtplib
import re
from colorama import init, Fore, Style
from datetime import datetime
from pathlib import Path
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from typing import List, Dict, Tuple, Optional

# Initialize colorama
init(autoreset=True)

# ============================================================================
# GITHUB SCANNER MODULE
# ============================================================================

class GitHubScanner:
    """Сканер GitHub репозиториев"""
    
    SUPPORTED_EXTENSIONS = {'.py', '.js', '.php', '.html', '.java', '.c', '.cpp', '.cs', '.go', '.rb', '.ts'}
    
    def parse_url(self, url: str) -> Optional[Dict]:
        """Парсит GitHub URL"""
        patterns = [
            r'github\.com/([^/]+)/([^/]+?)(?:/tree/([^/]+))?(?:\.git)?/?$',
            r'github\.com/([^/]+)/([^/]+?)/?$'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                owner, repo, branch = match.groups()
                repo = repo.replace('.git', '')
                return {'owner': owner, 'repo': repo, 'branch': branch or None}
        return None
    
    def get_files_to_scan(self, repo_info: Dict) -> Tuple[List[str], Optional[str]]:
        """Получает список файлов для сканирования"""
        owner, repo = repo_info['owner'], repo_info['repo']
        branch = repo_info.get('branch')
        
        try:
            # Получаем информацию о репозитории для определения ветки по умолчанию
            if not branch:
                resp = requests.get(f"https://api.github.com/repos/{owner}/{repo}", timeout=10)
                if resp.status_code == 200:
                    branch = resp.json().get('default_branch', 'main')
                else:
                    branch = 'main'
            
            # Получаем дерево файлов
            tree_url = f"https://api.github.com/repos/{owner}/{repo}/git/trees/{branch}?recursive=1"
            response = requests.get(tree_url, timeout=30)
            
            if response.status_code != 200:
                return [], None
            
            tree_data = response.json()
            files = []
            
            for item in tree_data.get('tree', []):
                if item['type'] == 'blob':
                    file_path = item['path']
                    if any(file_path.endswith(ext) for ext in self.SUPPORTED_EXTENSIONS):
                        files.append(file_path)
            
            return files, branch
        except Exception as e:
            print(f"{Fore.RED}[!] Ошибка при получении файлов GitHub: {e}")
            return [], None
    
    def read_file_content(self, repo_info: Dict, file_path: str) -> Optional[str]:
        """Читает содержимое файла из GitHub"""
        try:
            owner, repo, branch = repo_info['owner'], repo_info['repo'], repo_info.get('branch', 'main')
            url = f"https://raw.githubusercontent.com/{owner}/{repo}/{branch}/{file_path}"
            response = requests.get(url, timeout=10)
            
            if response.status_code == 200:
                return response.text
            return None
        except Exception:
            return None


# ============================================================================
# FILE SCANNER MODULE
# ============================================================================

class FileScanner:
    """Сканер локальных файлов"""
    
    SUPPORTED_EXTENSIONS = {'.py', '.js', '.php', '.html', '.java', '.c', '.cpp', '.cs', '.go', '.rb', '.ts'}
    
    def __init__(self, directory: str):
        self.directory = directory
    
    def get_files_to_scan(self) -> List[str]:
        """Получает список файлов для сканирования"""
        files = []
        try:
            for root, dirs, filenames in os.walk(self.directory):
                for filename in filenames:
                    if any(filename.endswith(ext) for ext in self.SUPPORTED_EXTENSIONS):
                        files.append(os.path.join(root, filename))
        except Exception as e:
            print(f"{Fore.RED}[!] Ошибка при сканировании папки: {e}")
        return files
    
    def read_file_content(self, file_path: str) -> Optional[str]:
        """Читает содержимое файла"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            return None


# ============================================================================
# AI SCANNER MODULE
# ============================================================================

class AIScanner:
    """Анализатор кода на основе AI"""
    
    def __init__(self):
        self.model = "gpt-4"
    
    async def analyze_code(self, file_path: str, code_content: str) -> List[Dict]:
        """Анализирует код на уязвимости"""
        try:
            import g4f

            prompt = f"""
                    ### РОЛЬ
                    Действуй как Старший инженер по безопасности приложений (Senior Application Security Engineer) и Аудитор кода. Твоя задача — провести глубокий статический анализ предоставленного исходного кода для выявления уязвимостей безопасности, логических ошибок и небезопасных практик программирования.

                    ### КОНТЕКСТ
                    Файл: {file_path}
                    Язык: Определи исходя из содержимого кода.

                    ### РУКОВОДСТВО ПО АНАЛИЗУ
                    1. **Отслеживание потока данных (Data Flow Tracking):** Отслеживай пользовательские входные данные (источники) до чувствительных операций (стоки). Не просто ищи ключевые слова; проверяй, действительно ли санитизация или валидация эффективны.
                    2. **Область угроз:** Ищи уязвимости OWASP Top 10 (SQLi, XSS, RCE, SSRF, IDOR), но ТАКЖЕ проверяй:
                       - Захардкоженные секреты (API-ключи, пароли, токены).
                       - Раскрытие информации (подробные сообщения об ошибках, стек трейсы).
                       - Небезопасные конфигурации (режим отладки включен, слабый CORS).
                       - Логические ошибки (состояния гонки, обходимые проверки).
                       - Устаревшие или небезопасные функции/библиотеки.
                    3. **Скрытые дефекты:** Обращай внимание на неочевидные уязвимости (oversights), такие как отсутствие проверки длины ввода, неправильное приведение типов или потенциальные уязвимости при расширении кода.

                    МАКСИМАЛЬНО ИЩИ КИБЕР УГРОЗЫ В КОДЕ!!!!!!!

                    ### УРОВНИ СЕРЬЕЗНОСТИ
                    - **Высокая:** Эксплуатируемо напрямую без сложных условий (например, незаэкранированный SQL-запрос, удаленное выполнение кода, обход аутентификации).
                    - **Средняя:** Эксплуатируемо при определенных условиях или требует взаимодействия пользователя (например, stored XSS, CSRF без токена, небезопасные прямые ссылки на объекты).
                    - **Незначительная:** Нарушения лучших практик, мелкие утечки информации или теоретические риски с низким воздействием (например, отсутствующие заголовки безопасности, подробное логирование несensitive данных).

                    ### ФОРМАТ ВЫВОДА
                    Верни результат **ТОЛЬКО** в виде валидного JSON-списка объектов. Не включай markdown-форматирование (как ```json), объяснения или текст вне JSON.
                    Если угроз не найдено, верни пустой список [].

                    МАКСИМАЛЬНО ИЩИ КИБЕР УГРОЗЫ В КОДЕ!!!!!!!

                    Каждый объект должен строго следовать этой схеме:
                    {{
                      "type": "Конкретное название уязвимости (например, SQL Injection, Hardcoded Secret)",
                      "severity": "High" | "Medium" | "Low",
                      "description": "Подробное объяснение, ПОЧЕМУ это ошибка, со ссылками на конкретные переменные или логику строк.",
                      "recommendation": "Конкретное исправление на уровне кода или рекомендация библиотеки."
                    }}

                    !!! Максимально подробно расписывай рекомендации. Текста должно быть много! Вставляй ссылки на документации по исправлению именно этой ошибки.

                    ВАЖНО: Все текстовые значения внутри JSON (поля type, description, recommendation) должны быть написаны на РУССКОМ языке. Значения поля severity оставь на английском (High, Medium, Low).

                    МАКСИМАЛЬНО ИЩИ КИБЕР УГРОЗЫ В КОДЕ!!!!!!!
                    ### КОД ДЛЯ АНАЛИЗА
                    {code_content}
                """
            
            response = await g4f.ChatCompletion.create_async(
                model=g4f.models.gpt_4,
                messages=[{"role": "user", "content": prompt}]
            )
            
            # Парсим JSON ответ
            try:
                result = json.loads(response)
                if isinstance(result, list):
                    return result
                elif isinstance(result, dict) and 'vulnerabilities' in result:
                    return result['vulnerabilities']
            except json.JSONDecodeError:
                # Пытаемся извлечь JSON из текста
                json_match = re.search(r'\[.*\]', response, re.DOTALL)
                if json_match:
                    return json.loads(json_match.group())
            
            return []
        except Exception as e:
            return []


# ============================================================================
# REPORT GENERATOR MODULE (python-docx + docx2pdf)
# ============================================================================

class ReportGenerator:
    """Генератор отчетов на базе python-docx + docx2pdf"""
    
    def __init__(self):
        pass
    
    def _add_styled_paragraph(self, doc, text: str, style_name: str = 'Normal', bold: bool = False, 
                              italic: bool = False, color: str = None, align: str = None):
        """Вспомогательный метод для добавления стилизованного параграфа"""
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        p = doc.add_paragraph(text, style=style_name)
        run = p.runs[0]
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if align:
            p.alignment = {
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'right': WD_ALIGN_PARAGRAPH.RIGHT
            }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        return p
    
    def generate_pdf_report(self, vulnerabilities: List[Dict], source: str, 
                          is_github: bool, filename: str = "report.pdf") -> str:
        """Генерирует PDF отчет через DOCX + конвертацию"""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx2pdf import convert
        
        doc = Document()
        
        # Настройка стилей для кириллицы
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Заголовок
        title = doc.add_heading('🛡️ WebGuardAI - Отчет о безопасности', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(44, 62, 80)
        title.runs[0].font.bold = True
        
        # Мета-информация
        doc.add_paragraph('')  # отступ
        meta = doc.add_paragraph()
        meta.add_run(f'📅 Дата: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        meta.add_run(f'📁 Источник: {source}\n')
        meta.add_run(f'🔗 Тип: {"GitHub репозиторий" if is_github else "Локальная папка"}')
        
        # Блок с итогами
        doc.add_heading('📊 Итоги', level=2)
        high = sum(1 for v in vulnerabilities if "high" in v.get('severity', '').lower())
        medium = sum(1 for v in vulnerabilities if "medium" in v.get('severity', '').lower())
        low = sum(1 for v in vulnerabilities if "low" in v.get('severity', '').lower())
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        summary_data = [
            ('Всего уязвимостей:', str(len(vulnerabilities))),
            ('🔴 Критичные (High):', str(high)),
            ('🟡 Средние (Medium):', str(medium)),
            ('🟢 Низкие (Low):', str(low)),
        ]
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
            # Подсветка критичных
            if "High" in label and high > 0:
                summary_table.cell(i, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                summary_table.cell(i, 1).paragraphs[0].runs[0].font.bold = True
        
        # Детальные результаты
        doc.add_heading('🔍 Детальные результаты', level=2)
        
        if vulnerabilities:
            for i, vuln in enumerate(vulnerabilities, 1):
                # Заголовок уязвимости
                severity = vuln.get('severity', 'Unknown').upper()
                sev_color = {'HIGH': 'FF0000', 'MEDIUM': 'FFAA00', 'LOW': '00AA00'}.get(severity, '666666')
                
                p = doc.add_paragraph()
                run = p.add_run(f'{i}. [{severity}] {vuln.get("type", "Unknown")}')
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(sev_color)
                run.font.size = Pt(11)
                
                # Информация
                doc.add_paragraph(f'📄 Файл: {vuln.get("file", "N/A")}', style='List Bullet')
                
                # Описание
                p_desc = doc.add_paragraph('📝 Описание: ', style='List Bullet')
                p_desc.add_run(vuln.get('description', 'N/A')).italic = True
                
                # Рекомендация
                p_rec = doc.add_paragraph('💡 Рекомендация: ', style='List Bullet')
                p_rec.add_run(vuln.get('recommendation', 'N/A'))
                
                doc.add_paragraph('─' * 60)  # разделитель
        else:
            doc.add_paragraph('✅ Уязвимостей не обнаружено.', style='Intense Quote')
        
        # Футер
        doc.add_paragraph('')
        footer = doc.add_paragraph('Отчет создан WebGuardAI - Advanced Code Security Scanner')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.italic = True
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        # Сохраняем временный DOCX
        temp_docx = filename.replace('.pdf', '.docx') if filename.endswith('.pdf') else filename + '.docx'
        doc.save(temp_docx)
        
        # Конвертируем в PDF
        try:
            convert(temp_docx, filename)
            # Удаляем временный DOCX
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
        except Exception as e:
            print(f"{Fore.YELLOW}[!] Не удалось конвертировать в PDF: {e}")
            print(f"{Fore.YELLOW}[!] Отчет сохранен в формате DOCX: {temp_docx}")
            return temp_docx
        
        return filename
    
    def generate_txt_report(self, vulnerabilities: List[Dict], source: str, 
                          is_github: bool, filename: str = "report.txt") -> str:
        """Генерирует текстовый отчет"""
        with open(filename, "w", encoding="utf-8") as f:
            f.write("=" * 70 + "\n")
            f.write("WebGuardAI - Отчет о безопасности\n")
            f.write("=" * 70 + "\n\n")
            
            f.write(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Источник: {source}\n")
            f.write(f"Тип: {'GitHub репозиторий' if is_github else 'Локальная папка'}\n\n")
            
            f.write("ИТОГИ\n")
            f.write("-" * 70 + "\n")
            high = sum(1 for v in vulnerabilities if "high" in v.get('severity', '').lower())
            medium = sum(1 for v in vulnerabilities if "medium" in v.get('severity', '').lower())
            low = sum(1 for v in vulnerabilities if "low" in v.get('severity', '').lower())
            
            f.write(f"Всего уязвимостей: {len(vulnerabilities)}\n")
            f.write(f"Критичные: {high}\n")
            f.write(f"Средние: {medium}\n")
            f.write(f"Низкие: {low}\n\n")
            
            f.write("ДЕТАЛЬНЫЕ РЕЗУЛЬТАТЫ\n")
            f.write("=" * 70 + "\n\n")
            
            for i, vuln in enumerate(vulnerabilities, 1):
                f.write(f"{i}. [{vuln.get('severity', 'Unknown').upper()}] {vuln.get('type', 'Unknown')}\n")
                f.write(f"   Файл: {vuln.get('file', 'N/A')}\n")
                f.write(f"   Описание: {vuln.get('description', 'N/A')}\n")
                f.write(f"   Рекомендация: {vuln.get('recommendation', 'N/A')}\n")
                f.write("-" * 70 + "\n\n")
        
        return filename


# ============================================================================
# MAIL SENDER MODULE
# ============================================================================

class MailSender:
    """Отправитель отчетов по почте"""
    
    def __init__(self, smtp_server: str, smtp_port: int, 
                 sender_email = "", sender_password = ""):
        self.smtp_server = smtp_server
        self.smtp_port = smtp_port
        self.sender_email = sender_email
        self.sender_password = sender_password
    
    def send_report(self, recipient_email: str, subject: str, vulnerabilities: List[Dict],
                   source: str, is_github: bool, attachment_path: Optional[str] = None,
                   attachment_type: str = "txt") -> Tuple[bool, str]:
        """Отправляет отчет по почте"""
        try:
            message = MIMEMultipart()
            message["From"] = self.sender_email or "WebGuardAI"
            message["To"] = recipient_email
            message["Subject"] = subject
            
            body = self._generate_email_body(vulnerabilities, source, is_github)
            message.attach(MIMEText(body, "plain"))
            
            if attachment_path and os.path.exists(attachment_path):
                self._attach_file(message, attachment_path, attachment_type)
            
            with smtplib.SMTP(self.smtp_server, self.smtp_port) as server:
                server.starttls()
                if self.sender_email and self.sender_password:
                    server.login(self.sender_email, self.sender_password)
                server.send_message(message)
            
            return True, "Письмо отправлено успешно"
        
        except smtplib.SMTPAuthenticationError:
            return False, "Ошибка аутентификации SMTP. Проверьте учетные данные."
        except smtplib.SMTPException as e:
            return False, f"Ошибка SMTP: {str(e)}"
        except Exception as e:
            return False, f"Ошибка отправки письма: {str(e)}"
    
    def _generate_email_body(self, vulnerabilities: List[Dict], source: str, is_github: bool) -> str:
        """Генерирует тело письма"""
        high = sum(1 for v in vulnerabilities if "high" in v.get('severity', '').lower())
        medium = sum(1 for v in vulnerabilities if "medium" in v.get('severity', '').lower())
        low = sum(1 for v in vulnerabilities if "low" in v.get('severity', '').lower())
        
        body = f"""
WebGuardAI - Отчет о сканировании безопасности
{'='*50}

Детали сканирования:
- Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- Источник: {source}
- Тип: {'GitHub репозиторий' if is_github else 'Локальная папка'}

Итоги:
- Всего уязвимостей: {len(vulnerabilities)}
- Критичные: {high}
- Средние: {medium}
- Низкие: {low}

Топ находок:
"""
        
        for i, vuln in enumerate(vulnerabilities[:5], 1):
            body += f"\n{i}. [{vuln.get('severity', 'Unknown').upper()}] {vuln.get('type', 'Unknown')}\n"
            body += f"   Файл: {vuln.get('file', 'N/A')}\n"
            body += f"   Описание: {vuln.get('description', 'N/A')[:100]}...\n"
        
        if len(vulnerabilities) > 5:
            body += f"\n... и еще {len(vulnerabilities) - 5} уязвимостей\n"
        
        body += f"""
{'='*50}
Подробную информацию смотрите в приложенном отчете.

Это письмо создано WebGuardAI - Advanced Code Security Scanner
"""
        return body
    
    def _attach_file(self, message, file_path: str, file_type: str = "txt"):
        """Прикрепляет файл к письму"""
        try:
            with open(file_path, "rb") as attachment:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(attachment.read())
            
            encoders.encode_base64(part)
            filename = os.path.basename(file_path)
            part.add_header("Content-Disposition", f'attachment; filename="{filename}"')
            message.attach(part)
        except Exception as e:
            print(f"Ошибка при прикреплении файла: {str(e)}")


# ============================================================================
# CLI MODULE
# ============================================================================

class WebGuardCLI:
    """Командная строка WebGuardAI"""
    
    def __init__(self):
        self.ai_scanner = AIScanner()
        self.github_scanner = GitHubScanner()
        self.report_generator = ReportGenerator()
    
    def print_banner(self):
        """Выводит баннер"""
        banner = pyfiglet.figlet_format("WebGuardAI", font="big")
        print(f"{Fore.CYAN}{Style.BRIGHT}{banner}")
        print(f"{Fore.WHITE}{'='*60}")
        print(f"{Fore.GREEN} WebGuardAI v3.0 - Сканер безопасности кода на базе AI")
        print(f"{Fore.WHITE}{'='*60}\n")
    
    def print_config(self, source: str, is_github: bool):
        """Выводит конфигурацию"""
        print(f"{Fore.CYAN}[*] Конфигурация:")
        print(f"    {Fore.WHITE}Тип источника: {'GitHub репозиторий' if is_github else 'Локальная папка'}")
        print(f"    {Fore.WHITE}Цель:          {source}")
        print(f"    {Fore.WHITE}Время:         {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print(f"    {Fore.WHITE}Расширения:    .py, .js, .php, .html, .java, .c, .cpp, .cs, .go, .rb, .ts")
        print(f"\n{Fore.WHITE}{'='*60}\n")
    
    def print_result_row(self, severity: str, vuln_type: str, file_path: str):
        """Выводит строку результата"""
        color = Fore.WHITE
        sev_lower = severity.lower()
        if "high" in sev_lower:
            color = Fore.RED
        elif "medium" in sev_lower:
            color = Fore.YELLOW
        elif "low" in sev_lower:
            color = Fore.GREEN
        
        print(f"{Fore.WHITE}[{color}{severity.upper():<7}{Fore.WHITE}] "
              f"{Fore.CYAN}{vuln_type:<20} "
              f"{Fore.WHITE}Файл: {Fore.BLUE}{file_path}")
    
    async def run_scan(self, source: str, is_github: bool, pdf_output: Optional[str] = None,
                      email_config: Optional[Dict] = None):
        """Запускает сканирование"""
        self.print_banner()
        self.print_config(source, is_github)
        
        if is_github:
            repo_info = self.github_scanner.parse_url(source)
            if not repo_info:
                print(f"{Fore.RED}[!] Ошибка: Неверный URL GitHub.")
                return
            
            print(f"{Fore.YELLOW}[*] Получение списка файлов из GitHub...")
            files, branch = self.github_scanner.get_files_to_scan(repo_info)
            if branch:
                repo_info['branch'] = branch
        else:
            file_scanner = FileScanner(source)
            files = file_scanner.get_files_to_scan()
        
        total = len(files)
        if total == 0:
            print(f"{Fore.RED}[!] Поддерживаемые файлы не найдены.")
            return
        
        print(f"{Fore.GREEN}[+] Найдено {total} файлов. Начинаем анализ AI...\n")
        print(f"{Fore.WHITE}{'КРИТИЧНОСТЬ':<12} {'ТИП':<21} {'ФАЙЛ'}")
        print(f"{Fore.WHITE}{'-'*60}")
        
        all_vulnerabilities = []
        
        for i, file_path in enumerate(files):
            filename = os.path.basename(file_path)
            sys.stdout.write(f"\r{Fore.YELLOW}[*] Прогресс: {i+1}/{total} - Сканирование: {filename[:20]:<20}")
            sys.stdout.flush()
            
            if is_github:
                content = self.github_scanner.read_file_content(repo_info, file_path)
            else:
                file_scanner = FileScanner(source)
                content = file_scanner.read_file_content(file_path)
            
            if content and content.strip():
                try:
                    vulnerabilities = await self.ai_scanner.analyze_code(file_path, content)
                    if vulnerabilities:
                        sys.stdout.write("\r" + " " * 70 + "\r")
                        for v in vulnerabilities:
                            v['file'] = file_path
                            self.print_result_row(v.get('severity', 'Low'), v.get('type', 'Unknown'), file_path)
                            all_vulnerabilities.append(v)
                except Exception:
                    pass
        
        print(f"\r{Fore.GREEN}[+] Сканирование завершено!{' '*50}")
        print(f"\n{Fore.WHITE}{'='*60}")
        print(f"{Fore.CYAN}[*] Итоги:")
        print(f"    {Fore.WHITE}Всего файлов:    {total}")
        print(f"    {Fore.WHITE}Всего угроз:     {len(all_vulnerabilities)}")
        
        high = sum(1 for v in all_vulnerabilities if "high" in v.get('severity', '').lower())
        med = sum(1 for v in all_vulnerabilities if "medium" in v.get('severity', '').lower())
        low = sum(1 for v in all_vulnerabilities if "low" in v.get('severity', '').lower())
        
        print(f"    {Fore.RED}Критичные:       {high}")
        print(f"    {Fore.YELLOW}Средние:         {med}")
        print(f"    {Fore.GREEN}Низкие:          {low}")
        print(f"{Fore.WHITE}{'='*60}\n")
        
        if all_vulnerabilities:
            self._handle_report_generation(all_vulnerabilities, source, is_github, pdf_output, email_config)
    
    def _handle_report_generation(self, vulnerabilities: List[Dict], source: str, 
                                 is_github: bool, pdf_output: Optional[str], 
                                 email_config: Optional[Dict]):
        """Обрабатывает генерацию отчетов"""
        txt_filename = "report.txt"
        self.report_generator.generate_txt_report(vulnerabilities, source, is_github, txt_filename)
        print(f"{Fore.GREEN}[+] Текстовый отчет сохранен в {txt_filename}")
        
        pdf_filename = None
        if pdf_output:
            pdf_filename = pdf_output if pdf_output.endswith('.pdf') else f"{pdf_output}.pdf"
            self.report_generator.generate_pdf_report(vulnerabilities, source, is_github, pdf_filename)
            print(f"{Fore.GREEN}[+] PDF отчет сохранен в {pdf_filename}")
        
        if email_config:
            self._send_email_report(vulnerabilities, source, is_github, email_config, 
                                   pdf_filename or txt_filename)
    
    def _send_email_report(self, vulnerabilities: List[Dict], source: str, is_github: bool,
                          email_config: Dict, attachment_path: str):
        """Отправляет отчет по почте"""
        try:
            print(f"{Fore.YELLOW}[*] Отправка отчета по почте...")
            
            mail_sender = MailSender(
                smtp_server=email_config['smtp_server'],
                smtp_port=email_config['smtp_port'],
                sender_email=email_config.get('sender_email'),
                sender_password=email_config.get('sender_password')
            )
            
            subject = f"WebGuardAI - Отчет о безопасности {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            success, message = mail_sender.send_report(
                recipient_email=email_config['recipient_email'],
                subject=subject,
                vulnerabilities=vulnerabilities,
                source=source,
                is_github=is_github,
                attachment_path=attachment_path,
                attachment_type="pdf" if attachment_path.endswith('.pdf') else "txt"
            )
            
            if success:
                print(f"{Fore.GREEN}[+] {message}")
            else:
                print(f"{Fore.RED}[!] {message}")
        
        except Exception as e:
            print(f"{Fore.RED}[!] Ошибка отправки письма: {str(e)}")


# ============================================================================
# GUI MODULE
# ============================================================================

def launch_gui():
    """Запускает графический интерфейс"""
    try:
        from PyQt6.QtWidgets import (
            QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
            QPushButton, QFileDialog, QTableWidget, QTableWidgetItem,
            QLabel, QHeaderView, QMessageBox, QProgressBar, QTextEdit, QGroupBox, QGridLayout,
            QLineEdit
        )
        from PyQt6.QtCore import Qt, QThread, pyqtSignal
        
        class ScanWorkerGUI(QThread):
            progress_signal = pyqtSignal(int)
            result_signal = pyqtSignal(list)
            finished_signal = pyqtSignal()
            status_signal = pyqtSignal(str)
            
            def __init__(self, source, is_github=False):
                super().__init__()
                self.source = source
                self.is_github = is_github
                self.is_running = True
            
            def run(self):
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                
                ai_scanner = AIScanner()
                
                if self.is_github:
                    github_scanner = GitHubScanner()
                    repo_info = github_scanner.parse_url(self.source)
                    if not repo_info:
                        self.status_signal.emit("Неверный URL GitHub")
                        self.finished_signal.emit()
                        return
                    
                    self.status_signal.emit(f"Получение файлов из GitHub: {repo_info['repo']}")
                    files, detected_branch = github_scanner.get_files_to_scan(repo_info)
                    if detected_branch:
                        repo_info['branch'] = detected_branch
                else:
                    file_scanner = FileScanner(self.source)
                    files = file_scanner.get_files_to_scan()
                
                total_files = len(files)
                
                if total_files == 0:
                    self.status_signal.emit("Файлы не найдены.")
                    self.finished_signal.emit()
                    return
                
                all_vulnerabilities = []
                
                for i, file_path in enumerate(files):
                    if not self.is_running:
                        break
                    
                    self.status_signal.emit(f"Сканирование: {os.path.basename(file_path)}")
                    
                    if self.is_github:
                        content = github_scanner.read_file_content(repo_info, file_path)
                    else:
                        content = file_scanner.read_file_content(file_path)
                    
                    if content and content.strip():
                        try:
                            vulnerabilities = loop.run_until_complete(ai_scanner.analyze_code(file_path, content))
                            if vulnerabilities:
                                for v in vulnerabilities:
                                    v['file'] = file_path
                                    all_vulnerabilities.append(v)
                        except Exception:
                            pass
                    
                    progress = int(((i + 1) / total_files) * 100)
                    self.progress_signal.emit(progress)
                
                loop.close()
                self.result_signal.emit(all_vulnerabilities)
                self.finished_signal.emit()
            
            def stop(self):
                self.is_running = False
        
        class MainWindowGUI(QMainWindow):
            def __init__(self):
                super().__init__()
                self.setWindowTitle("WebGuardAI")
                self.resize(1000, 700)
                
                self.central_widget = QWidget()
                self.setCentralWidget(self.central_widget)
                self.main_layout = QVBoxLayout(self.central_widget)
                self.main_layout.setSpacing(10)
                self.main_layout.setContentsMargins(10, 10, 10, 10)
                
                # Source section
                self.source_group = QGroupBox("Источник анализа")
                self.source_layout = QGridLayout()
                self.source_layout.setSpacing(10)
                
                self.lbl_local_folder = QLabel("Локальная папка:")
                self.input_local_folder = QLineEdit()
                self.input_local_folder.setReadOnly(True)
                self.btn_browse = QPushButton("Обзор")
                self.btn_browse.setStyleSheet("""
                    QPushButton {
                        background-color: #2196F3;
                        color: white;
                        border: none;
                        padding: 5px 15px;
                        font-weight: bold;
                    }
                    QPushButton:hover { background-color: #1976D2; }
                """)
                self.btn_browse.clicked.connect(self.select_folder)
                
                self.source_layout.addWidget(self.lbl_local_folder, 0, 0)
                self.source_layout.addWidget(self.input_local_folder, 0, 1, 1, 2)
                self.source_layout.addWidget(self.btn_browse, 0, 3)
                
                self.lbl_github = QLabel("GitHub URL:")
                self.input_github = QLineEdit()
                self.input_github.setPlaceholderText("https://github.com/...")
                self.input_github.textChanged.connect(self.on_github_url_changed)
                
                self.source_layout.addWidget(self.lbl_github, 1, 0)
                self.source_layout.addWidget(self.input_github, 1, 1, 1, 3)
                
                self.source_group.setLayout(self.source_layout)
                self.main_layout.addWidget(self.source_group)
                
                # Controls
                self.controls_layout = QHBoxLayout()
                
                self.btn_start_scan = QPushButton("Запустить сканирование")
                self.btn_start_scan.setStyleSheet("""
                    QPushButton {
                        background-color: #4CAF50;
                        color: white;
                        border: none;
                        padding: 8px 20px;
                        font-size: 14px;
                        font-weight: bold;
                        border-radius: 4px;
                    }
                    QPushButton:hover { background-color: #45a049; }
                    QPushButton:disabled { background-color: #cccccc; color: #666666; }
                """)
                self.btn_start_scan.clicked.connect(self.start_scan)
                self.btn_start_scan.setEnabled(False)
                
                self.progress_bar = QProgressBar()
                self.progress_bar.setTextVisible(False)
                self.progress_bar.setFixedHeight(25)
                self.progress_bar.setStyleSheet("""
                    QProgressBar {
                        border: 1px solid #ccc;
                        border-radius: 4px;
                        background-color: #f0f0f0;
                    }
                    QProgressBar::chunk {
                        background-color: #2196F3;
                        border-radius: 3px;
                    }
                """)
                
                self.lbl_status = QLabel("Готов к сканированию")
                self.lbl_status.setStyleSheet("color: #666; font-style: italic;")
                self.lbl_status.setMinimumWidth(150)
                self.lbl_status.setAlignment(Qt.AlignmentFlag.AlignRight)
                
                self.controls_layout.addWidget(self.btn_start_scan)
                self.controls_layout.addWidget(self.progress_bar, 1)
                self.controls_layout.addWidget(self.lbl_status)
                
                self.main_layout.addLayout(self.controls_layout)
                
                # Results
                self.results_group = QGroupBox("Результаты сканирования")
                self.results_layout = QVBoxLayout()
                
                self.table = QTableWidget()
                self.table.setColumnCount(4)
                headers = ["Критичность", "Тип уязвимости", "Файл", "Описание"]
                self.table.setHorizontalHeaderLabels(headers)
                
                header = self.table.horizontalHeader()
                header.setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
                header.setStyleSheet("""
                    QHeaderView::section {
                        background-color: #2196F3;
                        color: white;
                        padding: 8px;
                        border: none;
                        font-weight: bold;
                    }
                """)
                self.table.setStyleSheet("""
                    QTableWidget {
                        background-color: white;
                        color: black;
                        border: 1px solid #ccc;
                        gridline-color: #eee;
                    }
                    QTableWidget::item {
                        color: black;
                    }
                    QTableWidget::item:selected {
                        background-color: #e3f2fd;
                        color: black;
                    }
                """)
                self.table.cellClicked.connect(self.show_details)
                
                self.results_layout.addWidget(self.table)
                self.results_group.setLayout(self.results_layout)
                self.main_layout.addWidget(self.results_group)
                
                # Recommendations
                self.recommendations_group = QGroupBox("Рекомендации по исправлению")
                self.recommendations_layout = QVBoxLayout()
                
                self.details_text = QTextEdit()
                self.details_text.setReadOnly(True)
                self.details_text.setStyleSheet("""
                    QTextEdit {
                        background-color: white;
                        color: black;
                        border: 1px solid #ccc;
                        padding: 5px;
                        font-family: Consolas, monospace;
                    }
                """)
                
                self.recommendations_layout.addWidget(self.details_text)
                self.recommendations_group.setLayout(self.recommendations_layout)
                self.main_layout.addWidget(self.recommendations_group)
                
                self.target_dir = None
                self.vulnerabilities = []
            
            def on_github_url_changed(self):
                url = self.input_github.text().strip()
                if url:
                    self.btn_start_scan.setEnabled(True)
                    self.lbl_status.setText("URL GitHub введен")
                elif not self.target_dir:
                    self.btn_start_scan.setEnabled(False)
                    self.lbl_status.setText("Готов к сканированию")
            
            def select_folder(self):
                dir_path = QFileDialog.getExistingDirectory(self, "Выберите папку для сканирования")
                if dir_path:
                    self.target_dir = dir_path
                    self.input_local_folder.setText(dir_path)
                    self.input_github.clear()
                    self.btn_start_scan.setEnabled(True)
                    self.lbl_status.setText("Папка выбрана")
            
            def start_scan(self):
                github_url = self.input_github.text().strip()
                
                if github_url:
                    source = github_url
                    is_github = True
                elif self.target_dir:
                    source = self.target_dir
                    is_github = False
                else:
                    return
                
                self.table.setRowCount(0)
                self.vulnerabilities = []
                self.btn_start_scan.setEnabled(False)
                self.btn_browse.setEnabled(False)
                self.input_github.setEnabled(False)
                self.progress_bar.setValue(0)
                self.lbl_status.setText("Сканирование...")
                
                self.worker = ScanWorkerGUI(source, is_github)
                self.worker.progress_signal.connect(self.update_progress)
                self.worker.result_signal.connect(self.process_results)
                self.worker.status_signal.connect(self.update_status)
                self.worker.finished_signal.connect(self.scan_finished)
                self.worker.start()
            
            def update_progress(self, value):
                self.progress_bar.setValue(value)
            
            def update_status(self, text):
                self.lbl_status.setText(text)
            
            def process_results(self, results):
                self.vulnerabilities = results
                self.table.setRowCount(len(results))
                for i, v in enumerate(results):
                    severity_item = QTableWidgetItem(v.get('severity', 'Unknown'))
                    severity_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    self.table.setItem(i, 0, severity_item)
                    
                    self.table.setItem(i, 1, QTableWidgetItem(v.get('type', 'Unknown')))
                    self.table.setItem(i, 2, QTableWidgetItem(os.path.basename(v.get('file', 'Unknown'))))
                    self.table.setItem(i, 3, QTableWidgetItem(
                        v.get('description', '')[:50] + "..." if len(v.get('description', '')) > 50 else v.get('description', '')))
            
            def scan_finished(self):
                self.btn_start_scan.setEnabled(True)
                self.btn_browse.setEnabled(True)
                self.input_github.setEnabled(True)
                self.lbl_status.setText("Готов к сканированию")
                QMessageBox.information(self, "Сканирование завершено", f"Найдено {len(self.vulnerabilities)} потенциальных угроз.")
            
            def show_details(self, row, column):
                if row < len(self.vulnerabilities):
                    v = self.vulnerabilities[row]
                    details = f"<h3>Детали уязвимости</h3>"
                    details += f"<p><b>Файл:</b> {v.get('file')}</p>"
                    details += f"<p><b>Тип:</b> {v.get('type')}</p>"
                    details += f"<p><b>Критичность:</b> <span style='color:red; font-weight:bold;'>{v.get('severity')}</span></p>"
                    details += f"<hr>"
                    details += f"<p><b>Описание:</b><br>{v.get('description', 'Описание отсутствует')}</p>"
                    details += f"<p><b>Рекомендация:</b><br>{v.get('recommendation', 'Рекомендация отсутствует')}</p>"
                    self.details_text.setHtml(details)
        
        app = QApplication(sys.argv)
        window = MainWindowGUI()
        window.show()
        sys.exit(app.exec())
    
    except ImportError:
        print(f"{Fore.RED}[!] PyQt6 не установлен. Установите: pip install PyQt6")
        sys.exit(1)
    except Exception as e:
        print(f"{Fore.RED}[!] Ошибка запуска GUI: {str(e)}")
        sys.exit(1)


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="WebGuardAI - Сканер безопасности кода на базе AI",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры использования:
  Сканирование локальной папки:
    python webguard.py -d /path/to/project
  
  Сканирование GitHub репозитория:
    python webguard.py -u https://github.com/owner/repo
  
  Сканирование и экспорт в PDF:
    python webguard.py -d /path/to/project --pdf report.pdf
  
  Сканирование и отправка по почте:
    python webguard.py -u https://github.com/owner/repo --email recipient@example.com --smtp-server smtp.gmail.com --smtp-port 587 --sender-email your@gmail.com --sender-password "your-password"
  
  Запуск GUI:
    python webguard.py --gui
        """
    )
    
    mode_group = parser.add_argument_group('Режим работы')
    mode_exclusive = mode_group.add_mutually_exclusive_group(required=False)
    mode_exclusive.add_argument("-d", "--dir", help="Локальная папка для сканирования")
    mode_exclusive.add_argument("-u", "--url", help="URL GitHub репозитория для сканирования")
    mode_exclusive.add_argument("--gui", action="store_true", help="Запустить графический интерфейс")
    
    report_group = parser.add_argument_group('Опции отчета')
    report_group.add_argument("--pdf", metavar="FILE", help="Сохранить отчет в PDF")
    report_group.add_argument("--txt", metavar="FILE", help="Сохранить отчет в TXT (по умолчанию: report.txt)")
    
    email_group = parser.add_argument_group('Опции почты')
    email_group.add_argument("--email", metavar="RECIPIENT", help="Отправить отчет на адрес электронной почты")
    email_group.add_argument("--smtp-server", default="smtp.hoster.by", help="SMTP сервер")
    email_group.add_argument("--smtp-port", type=int, default=465, help="Порт SMTP")
    email_group.add_argument("--sender-email", help="Email отправителя (опционально)")
    email_group.add_argument("--sender-password", help="Пароль email отправителя (опционально)")
    
    args = parser.parse_args()
    
    if args.gui:
        launch_gui()
        return
    
    if not args.dir and not args.url:
        parser.print_help()
        sys.exit(1)
    
    source = args.url if args.url else args.dir
    is_github = True if args.url else False
    
    email_config = None
    if args.email:
        email_config = {
            'recipient_email': args.email,
            'sender_email': args.sender_email,
            'sender_password': args.sender_password,
            'smtp_server': args.smtp_server,
            'smtp_port': args.smtp_port
        }
    
    cli = WebGuardCLI()
    
    try:
        asyncio.run(cli.run_scan(source, is_github, pdf_output=args.pdf, email_config=email_config))
    except KeyboardInterrupt:
        print(f"\n{Fore.RED}[!] Сканирование прервано пользователем.")
        sys.exit(0)


if __name__ == "__main__":
    main()